from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BLUE = "2F75B5"
DARK_BLUE = "1F4E78"
LIGHT_BLUE = "EAF2F8"
GRAY = "6B7280"
WHITE = "FFFFFF"


def shade_cell(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths=None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        shade_cell(cell, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
            run.font.size = Pt(9.5)

    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for column_index, value in enumerate(values):
            cells[column_index].text = str(value)
            cells[column_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[column_index])
            if row_index % 2 == 1:
                shade_cell(cells[column_index], "F4F7FA")
            for run in cells[column_index].paragraphs[0].runs:
                run.font.size = Pt(9.2)
        if widths:
            for column_index, width in enumerate(widths):
                cells[column_index].width = width
    document.add_paragraph()


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(4)
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor.from_string(DARK_BLUE if level == 1 else BLUE)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)
        paragraph.paragraph_format.space_after = Pt(2)


def add_callout(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.text = text
    shade_cell(cell, LIGHT_BLUE)
    set_cell_margins(cell, 120, 140, 120, 140)
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    document.add_paragraph()


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    for name, size in (("Title", 30), ("Heading 1", 18), ("Heading 2", 13)):
        styles[name].font.name = "Aptos Display"
        styles[name].font.size = Pt(size)

    header = section.header.paragraphs[0]
    header.text = "SUMLINK ANALYTICS DASHBOARD | MANAGER SUBMISSION"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(GRAY)

    footer = section.footer.paragraphs[0]
    footer.text = "Prepared by Harsha S | 17 August 2026"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(GRAY)


def add_cover(document: Document) -> None:
    document.add_paragraph().paragraph_format.space_after = Pt(55)
    brand = document.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = brand.add_run("SUMLINK")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(5)
    run = title.add_run("Sumlink Analytics Dashboard")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run("Production Implementation and Validation Report")
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(GRAY)

    add_table(
        document,
        ["Report item", "Details"],
        [
            ["Prepared for", "Management Review"],
            ["Prepared by", "Harsha S"],
            ["Report date", "17 August 2026"],
            ["GA4 property", "516899630"],
            ["Dashboard", "Appsmith Cloud"],
            ["Backend", "FastAPI on Render"],
        ],
        [Inches(1.8), Inches(4.6)],
    )
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "Confidentiality: Internal project report. Credentials, API keys, and service-account secrets are intentionally excluded."
    )
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    document.add_page_break()


def build_report(output_path: Path) -> None:
    document = Document()
    configure_document(document)
    add_cover(document)

    add_heading(document, "1. Executive Summary")
    add_body(document, "The Sumlink Analytics Dashboard is a production-ready product analytics solution for management and game-performance review. It combines an Appsmith Cloud interface with a FastAPI backend hosted on Render. The backend retrieves aggregated reporting data from the Google Analytics 4 Data API and exact user-level rolling-retention data from the GA4 BigQuery export.")
    add_body(document, "The delivered dashboard contains five management sections: Executive Health, Acquisition, Onboarding, Gameplay, and Retention. The production backend is live, GA4 connectivity is confirmed, exact rolling retention is enabled through BigQuery, and the focused backend validation suite passed all tests on 17 August 2026.")
    add_table(document, ["Area", "Status", "Evidence"], [
        ["Render backend", "Operational", "/health returned ok; version 1.0.0"],
        ["GA4 connection", "Connected", "/ready returned ready and GA4 connected"],
        ["Appsmith dashboard", "Deployed", "Five sections imported into Appsmith Cloud"],
        ["BigQuery retention", "Enabled", "Exact user-level rolling retention returned through the retention endpoint"],
        ["Automated tests", "Passed", "13 backend tests passed"],
    ])

    add_heading(document, "2. Project Objective")
    add_body(document, "The objective is to give management one consistent view of product health, user acquisition, onboarding quality, gameplay balance, and retention. The dashboard replaces fragmented manual checks with reusable filters, documented KPI definitions, and API-backed reporting.")
    add_bullets(document, [
        "Provide a concise management view of user growth, engagement, conversion, gameplay, and retention.",
        "Use live GA4 reporting data while keeping credentials outside the Appsmith client.",
        "Support exact rolling-retention analysis using BigQuery user-level events.",
        "Allow managers to review the deployed dashboard without running a developer laptop.",
    ])

    add_heading(document, "3. Production Architecture")
    add_table(document, ["Layer", "Technology", "Responsibility"], [
        ["Dashboard", "Appsmith Cloud", "Filters, KPI cards, charts, tables, and manager-facing navigation"],
        ["API", "Python FastAPI", "Authentication, validation, aggregation, caching, and response shaping"],
        ["Hosting", "Render", "Public HTTPS backend deployment and health monitoring"],
        ["Reporting data", "GA4 Data API", "Aggregated users, sessions, engagement, funnels, events, and cohort metrics"],
        ["Raw analytics", "GA4 BigQuery export", "Exact distinct-user rolling retention and user-level cohort calculation"],
    ])
    add_callout(document, "Data flow: Manager -> Appsmith Cloud -> HTTPS FastAPI endpoint -> GA4 Data API and/or BigQuery -> normalized JSON response -> dashboard cards, charts, and tables.")
    add_body(document, "Appsmith never receives the Google service-account credential. It sends only dashboard filters and the protected API header to the Render backend.")

    add_heading(document, "4. Hosting and Deployment")
    add_table(document, ["Component", "Production location", "Current state"], [
        ["Backend API", "https://sumlink-analytics-api.onrender.com", "Live and externally reachable"],
        ["Health check", "/health", "Status ok; service version 1.0.0"],
        ["Readiness check", "/ready", "Status ready; GA4 connected"],
        ["Dashboard", "xorstack.appsmith.com workspace", "Imported and deployed in Appsmith Cloud"],
        ["Local environment", "localhost", "Development copy only; not required for production access"],
    ])
    add_body(document, "Manager access must use the deployed Appsmith Cloud application URL. If login is required, invite the manager through Appsmith Share or make the application public according to company policy.")

    add_heading(document, "5. Dashboard Modules")
    modules = [
        ("5.1 Executive Health", ["Active users and new users for the selected period.", "DAU, MAU, and DAU/MAU stickiness.", "Engagement rate, sessions, screen views, and activity trends.", "A compact management view of product adoption and engagement."]),
        ("5.2 Acquisition", ["User acquisition and churn-oriented summary metrics.", "Channel, source, and campaign quality comparisons where GA4 dimensions are available.", "New-user growth and returning-user behavior for the selected filters."]),
        ("5.3 Onboarding", ["Onboarding and tutorial funnel progression.", "Tutorial started, completed, failed, skipped, and frustration indicators.", "Conversion and drop-off rates across onboarding steps.", "Average tutorial-step timing where registered GA4 parameters are available."]),
        ("5.4 Gameplay", ["Level starts, completions, failures, and drop-off analysis.", "New versus returning player activity.", "Hint and add-row usage, level difficulty, and completion trends.", "Level-level tables for identifying balancing and progression issues."]),
        ("5.5 Retention", ["DAU, WAU, MAU, stickiness, sessions, and returning users.", "Exact Day 1, Day 3, Day 7, and Day 15 cohort retention from GA4 reporting.", "Exact rolling Day 1+, Day 3+, Day 7+, and Day 15+ retention from BigQuery.", "Cohort tables and trends with mature-day handling for incomplete cohorts."]),
    ]
    for title, items in modules:
        add_heading(document, title, level=2)
        add_bullets(document, items)

    add_heading(document, "6. Global Filters")
    add_table(document, ["Filter", "Purpose"], [
        ["Date range", "Last 7, 14, or 30 days and supported custom date ranges"],
        ["App version", "Compare releases and identify version-specific changes"],
        ["OS version", "Segment platform behavior"],
        ["Device model", "Identify device-specific differences"],
        ["Country", "Review geographic performance"],
        ["Level", "Focus gameplay metrics on a selected level"],
    ])

    add_heading(document, "7. KPI Definitions and Formulas")
    add_table(document, ["KPI", "Definition or formula"], [
        ["Engagement rate", "Engaged sessions / Sessions x 100"],
        ["Stickiness", "DAU / MAU x 100"],
        ["Onboarding completion", "Users completing onboarding / Users starting onboarding x 100"],
        ["Level completion", "Completed level events / Started level events x 100"],
        ["Level drop-off", "(Level starts - Level completions) / Level starts x 100"],
        ["Failure rate", "Failed events / Started events x 100"],
        ["Skip rate", "Skipped events / Started events x 100"],
        ["Exact Day N retention", "Day 0 cohort users active exactly on Day N / Day 0 cohort users x 100"],
        ["Rolling Day N+ retention", "Day 0 cohort users active on Day N or any later day / Day 0 cohort users x 100"],
    ])

    add_heading(document, "8. GA4 and BigQuery Data Methodology")
    add_heading(document, "8.1 GA4 Data API", level=2)
    add_body(document, "The GA4 Data API supplies aggregated reporting including active users, new users, sessions, engaged sessions, event counts, registered custom dimensions, registered custom metrics, funnels, filters, and standard cohort-style metrics. It follows GA4 reporting definitions and privacy controls.")
    add_heading(document, "8.2 BigQuery GA4 Export", level=2)
    add_body(document, "BigQuery is used where exact user-level event history is required. The retention service scans the GA4 events_* export, derives each user's Day 0 cohort from first activity, calculates later activity offsets, and counts distinct user_pseudo_id values. This produces exact rolling retention rather than an aggregate estimate.")
    add_heading(document, "8.3 Rolling Retention", level=2)
    add_body(document, "Rolling Day N+ retention measures whether a user from a Day 0 cohort returned on Day N or any later day. A cohort day is displayed only when enough time has elapsed to mature. The All Users result is a weighted aggregate of eligible cohorts, preventing small cohorts from distorting the overall percentage.")

    add_heading(document, "9. Validation Results")
    add_body(document, "Validation timestamp: 17 August 2026 IST. These values were refreshed from the live Render service before this report was generated. Values can move as GA4 processes new events and as the selected dashboard date range changes.")
    add_table(document, ["Live measure", "Verified value"], [
        ["GA4 connection", "Connected"],
        ["Day 0 cohort users", "123"],
        ["Active users - selected range", "Live dashboard range dependent"],
        ["DAU / WAU / MAU", "24 / 125 / 396"],
        ["Stickiness", "6.06%"],
        ["Sessions / Engaged sessions", "1,752 / 1,315"],
        ["Returning users", "346"],
        ["Day 1 retention", "15.45%"],
        ["Day 3 retention", "2.44%"],
        ["Day 7 retention", "0.81%"],
        ["Day 15 retention", "0.00% (0 users)"],
        ["Rolling cohort rows", "Returned by live retention API"],
        ["Rolling data source", "BigQuery exact user-level"],
    ])
    add_table(document, ["Automated verification", "Result"], [
        ["BigQuery retention service tests", "Passed"],
        ["API response tests", "Passed"],
        ["Backend test total", "13 passed"],
        ["Render health check", "Passed"],
        ["Render readiness and GA4 connection", "Passed"],
    ])
    add_callout(document, "Validation conclusion: Standard Day N retention uses GA4 cohort reporting, while rolling Day N+ retention uses exact distinct-user calculations from the BigQuery export.")

    add_heading(document, "10. API Interface")
    add_table(document, ["Dashboard section", "Production API path"], [
        ["Executive Health", "/api/v1/executive/summary"],
        ["Acquisition", "/api/v1/acquisition/summary"],
        ["Onboarding", "/api/v1/onboarding/summary"],
        ["Gameplay", "/api/v1/gameplay/summary"],
        ["Retention", "/api/v1/retention/summary"],
    ])
    add_body(document, "Protected dashboard requests use the x-api-key header. The key value is intentionally not included in this report.")

    add_heading(document, "11. Security and Operational Controls")
    add_bullets(document, [
        "API authentication is enabled for protected production routes.",
        "Google service-account credentials are stored as backend environment secrets.",
        "The Appsmith interface does not contain or expose Google credentials.",
        "CORS is configured for the Appsmith Cloud origin.",
        "Backend responses are cached for 300 seconds to reduce repeated GA4 queries.",
        "Production readiness is monitored through dedicated health and readiness routes.",
    ])

    add_heading(document, "12. Limitations and Operational Notes")
    add_bullets(document, [
        "The Render free service can sleep after inactivity; the first request may require a short wake-up period.",
        "GA4 Data API results are aggregated and can differ from raw-event queries when definitions, identity, filters, time zones, or processing windows differ.",
        "BigQuery rolling retention depends on the availability and completeness of the GA4 events_* export.",
        "Recent cohorts correctly show unavailable values for retention days that have not yet matured.",
        "Manager access depends on Appsmith sharing configuration: public application or invited workspace user.",
        "Some historical dashboard comparison labels showed character-encoding artifacts; these do not change the validated KPI calculations.",
    ])

    add_heading(document, "13. Final Project Status")
    add_table(document, ["Deliverable", "Final status"], [
        ["Five-section analytics dashboard", "Complete"],
        ["FastAPI GA4 backend", "Deployed on Render"],
        ["Appsmith Cloud application", "Imported and deployed"],
        ["GA4 Data API integration", "Validated"],
        ["BigQuery exact rolling retention", "Validated"],
        ["Authentication and production configuration", "Enabled"],
        ["Manager documentation", "Complete"],
    ])

    add_heading(document, "14. 1-Year Analytics Roadmap")
    add_body(document, "The next phase is to move analytics from reactive reporting into a dependable decision-making system for product, growth, retention, monetization, and management review. The roadmap below uses weekly sprint execution so each improvement includes implementation, validation, documentation, and stakeholder review.")
    add_table(document, ["Quarter", "Focus", "Expected outcomes", "Key deliverables"], [
        ["Q1: Foundation and Data Reliability", "Stabilize tracking, dashboard accuracy, KPI logic, and reporting consistency.", "Clean GA4 event structure, reliable Appsmith and backend reporting, business-aligned KPI definitions, reduced manual checking.", "Analytics audit report; KPI dictionary; validated executive dashboard; data quality checklist; tracking issue backlog."],
        ["Q2: Funnel and User Journey Analytics", "Understand acquisition quality, onboarding flow, and drop-off points.", "Clear visibility from install to activation, funnel-level drop-off analysis, better identification of weak onboarding and conversion points.", "Acquisition dashboard; onboarding funnel dashboard; early behavior analysis; segment comparison views; funnel insights summary."],
        ["Q3: Retention, Cohorts, and Engagement", "Strengthen retention understanding using GA4 and BigQuery.", "Day 1, Day 3, Day 7, and Day 15+ retention visibility, rolling retention, cohort analysis, and engagement-retention links.", "Retention dashboard; BigQuery cohort validation; rolling retention model; engagement-retention correlation report; churn-risk indicators."],
        ["Q4: Automation, Optimization, and Decision Support", "Turn analytics into repeatable weekly and monthly decision support.", "Automated stakeholder reporting, faster product and growth decisions, experiment-readiness, and stronger governance.", "Automated reporting workflow; experiment analytics framework; executive health review pack; KPI cadence; annual analytics summary."],
    ])
    add_table(document, ["Sprint day", "Activity"], [
        ["Day 1", "Requirement alignment and sprint planning"],
        ["Day 2", "Data extraction, event review, and logic design"],
        ["Day 3", "Dashboard, query, or backend implementation"],
        ["Day 4", "Validation against GA4 and BigQuery"],
        ["Day 5", "Review, fixes, documentation, and stakeholder update"],
    ])
    add_table(document, ["Risk or dependency", "Management note"], [
        ["Stable GA4 event tracking", "Event names and parameters should be documented before schema changes."],
        ["BigQuery access and export continuity", "Exact rolling retention depends on continued GA4 BigQuery export availability."],
        ["KPI definition approval", "Business signoff prevents dashboard logic from drifting from management expectations."],
        ["Stakeholder availability", "Periodic review is needed to keep reports actionable and aligned."],
    ])

    add_heading(document, "15. Manager-Ready Conclusion")
    add_body(document, "The Sumlink Analytics Dashboard is ready for management review. It provides one consolidated view of executive health, acquisition, onboarding, gameplay, and retention. The production backend is live on Render, the dashboard is deployed in Appsmith Cloud, GA4 connectivity is confirmed, and exact rolling-retention calculations are supported by BigQuery user-level data.")
    add_body(document, "The manager should receive the deployed Appsmith Cloud application link rather than a localhost URL. Before sharing, confirm access in a private browser window and either enable approved public access or invite the manager's email address.")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


if __name__ == "__main__":
    destination = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("reports/Sumlink_Analytics_Dashboard_Final_Report.docx")
    build_report(destination.resolve())
    if len(sys.argv) > 2:
        try:
            import mammoth
        except ModuleNotFoundError as exc:
            raise SystemExit("HTML export requires the optional 'mammoth' package.") from exc
        html_destination = Path(sys.argv[2]).resolve()
        with destination.resolve().open("rb") as source:
            converted = mammoth.convert_to_html(source)
        html_document = f"""<!doctype html>
<html><head><meta charset=\"utf-8\"><title>Sumlink Analytics Dashboard</title>
<style>
@page {{ size: A4; margin: 16mm; }}
body {{ font-family: Arial, sans-serif; color: #1f2937; font-size: 10.5pt; line-height: 1.35; }}
h1 {{ color: #1f4e78; font-size: 21pt; border-bottom: 2px solid #2f75b5; padding-bottom: 4pt; margin-top: 18pt; }}
h2 {{ color: #2f75b5; font-size: 14pt; margin-top: 12pt; }}
p {{ margin: 5pt 0 8pt; }}
li {{ margin-bottom: 3pt; }}
table {{ width: 100%; border-collapse: collapse; margin: 8pt 0 14pt; break-inside: avoid; }}
td, th {{ border: 1px solid #cbd5e1; padding: 6pt; vertical-align: top; }}
tr:first-child td {{ color: white; background: #2f75b5; font-weight: bold; }}
tr:nth-child(even) td {{ background: #f4f7fa; }}
</style></head><body>{converted.value}</body></html>"""
        html_destination.write_text(html_document, encoding="utf-8")
    print(destination.resolve())
