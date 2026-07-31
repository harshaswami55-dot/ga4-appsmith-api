from __future__ import annotations

import csv
import json
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches
from google.analytics.data_v1beta import BetaAnalyticsDataClient
from google.oauth2 import service_account


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "backend"
sys.path.insert(0, str(BACKEND))

from app.config import get_settings  # noqa: E402


def cell(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float):
        return f"{value:.2f}".rstrip("0").rstrip(".")
    return str(value)


def field(obj: Any, name: str, default: Any = "") -> Any:
    return getattr(obj, name, default)


def rows_to_csv(path: Path, rows: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    keys: list[str] = []
    for row in rows:
        for key in row:
            if key not in keys:
                keys.append(key)
    with path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=keys or ["empty"])
        writer.writeheader()
        for row in rows:
            writer.writerow(row)


def run_report(
    client: BetaAnalyticsDataClient,
    property_id: str,
    *,
    dimensions: list[str],
    metrics: list[str],
    start_date: str = "90daysAgo",
    end_date: str = "yesterday",
    limit: int = 10000,
    order_metric: str | None = None,
) -> list[dict[str, Any]]:
    request: dict[str, Any] = {
        "property": f"properties/{property_id}",
        "date_ranges": [{"start_date": start_date, "end_date": end_date}],
        "dimensions": [{"name": d} for d in dimensions],
        "metrics": [{"name": m} for m in metrics],
        "limit": limit,
        "keep_empty_rows": False,
    }
    if order_metric:
        request["order_bys"] = [{"metric": {"metric_name": order_metric}, "desc": True}]
    response = client.run_report(request=request, timeout=60)
    dimension_names = [h.name for h in response.dimension_headers]
    metric_names = [h.name for h in response.metric_headers]
    output: list[dict[str, Any]] = []
    for row in response.rows:
        item: dict[str, Any] = {}
        for name, value in zip(dimension_names, row.dimension_values):
            item[name] = value.value
        for name, value in zip(metric_names, row.metric_values):
            raw = value.value
            try:
                item[name] = int(raw) if raw.isdigit() else float(raw)
            except Exception:
                item[name] = raw
        output.append(item)
    return output


def add_table(doc: Document, rows: list[dict[str, Any]], columns: list[str], max_rows: int = 40) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for i, col in enumerate(columns):
        table.rows[0].cells[i].text = col
    for row in rows[:max_rows]:
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            cells[i].text = cell(row.get(col, ""))
    if len(rows) > max_rows:
        doc.add_paragraph(f"Showing first {max_rows} of {len(rows)} rows. Full rows are in the CSV export.")


def main() -> None:
    settings = get_settings()
    property_id = settings.ga4_property_id
    credentials_path = Path(settings.google_application_credentials).expanduser()
    credentials = service_account.Credentials.from_service_account_file(
        credentials_path,
        scopes=["https://www.googleapis.com/auth/analytics.readonly"],
    )
    client = BetaAnalyticsDataClient(credentials=credentials)

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_dir = ROOT / "reports" / f"ga4_event_custom_parameter_audit_{stamp}"
    out_dir.mkdir(parents=True, exist_ok=True)

    metadata = client.get_metadata(name=f"properties/{property_id}/metadata")
    all_dimensions = [
        {
            "api_name": field(d, "api_name"),
            "ui_name": field(d, "ui_name"),
            "category": field(d, "category"),
            "description": field(d, "description"),
            "custom_definition": field(d, "custom_definition", False),
            "deprecated_api_names": ", ".join(field(d, "deprecated_api_names", [])),
        }
        for d in metadata.dimensions
    ]
    all_metrics = [
        {
            "api_name": field(m, "api_name"),
            "ui_name": field(m, "ui_name"),
            "category": field(m, "category"),
            "description": field(m, "description"),
            "custom_definition": field(m, "custom_definition", False),
            "type": field(m, "type_").name if hasattr(field(m, "type_"), "name") else field(m, "type_"),
            "deprecated_api_names": ", ".join(field(m, "deprecated_api_names", [])),
        }
        for m in metadata.metrics
    ]

    custom_dimensions = [
        d
        for d in all_dimensions
        if d["custom_definition"]
        or str(d["api_name"]).startswith(("customEvent:", "customUser:", "customItem:"))
    ]
    custom_metrics = [
        m
        for m in all_metrics
        if m["custom_definition"]
        or str(m["api_name"]).startswith(("customEvent:", "averageCustomEvent:", "customUser:", "customItem:"))
    ]

    events = run_report(
        client,
        property_id,
        dimensions=["eventName"],
        metrics=["eventCount", "totalUsers", "activeUsers"],
        order_metric="eventCount",
        limit=10000,
    )

    configured_definitions = [
        {"type": "dimension", "name": settings.level_dimension, "purpose": "Level filters and level charts"},
        {"type": "dimension", "name": settings.tutorial_step_dimension, "purpose": "Tutorial step funnel/drop-off"},
        {"type": "dimension", "name": settings.action_type_dimension, "purpose": "Gameplay action mapping such as add_row"},
        {"type": "metric", "name": settings.level_time_metric, "purpose": "Average time by level"},
        {"type": "metric", "name": settings.tutorial_time_metric, "purpose": "Average tutorial time by step"},
    ]
    validation_rows: list[dict[str, Any]] = []
    for item in configured_definitions:
        try:
            if item["type"] == "dimension":
                sample = run_report(
                    client,
                    property_id,
                    dimensions=[item["name"]],
                    metrics=["eventCount", "activeUsers"],
                    order_metric="eventCount",
                    limit=10,
                )
            else:
                sample = run_report(
                    client,
                    property_id,
                    dimensions=["eventName"],
                    metrics=[item["name"]],
                    order_metric=item["name"],
                    limit=10,
                )
            validation_rows.append(
                {
                    "type": item["type"],
                    "api_name": item["name"],
                    "purpose": item["purpose"],
                    "status": "Available in GA4 Data API",
                    "sample_rows": len(sample),
                    "top_sample": json.dumps(sample[:3], ensure_ascii=False),
                }
            )
        except Exception as exc:
            validation_rows.append(
                {
                    "type": item["type"],
                    "api_name": item["name"],
                    "purpose": item["purpose"],
                    "status": "Not available / query failed",
                    "sample_rows": 0,
                    "top_sample": str(exc)[:500],
                }
            )

    custom_dimension_samples: dict[str, list[dict[str, Any]]] = {}
    event_custom_samples: dict[str, list[dict[str, Any]]] = {}
    for d in custom_dimensions[:25]:
        api_name = d["api_name"]
        try:
            custom_dimension_samples[api_name] = run_report(
                client,
                property_id,
                dimensions=[api_name],
                metrics=["eventCount", "activeUsers"],
                order_metric="eventCount",
                limit=100,
            )
        except Exception as exc:
            custom_dimension_samples[api_name] = [{"error": str(exc)}]
        try:
            event_custom_samples[api_name] = run_report(
                client,
                property_id,
                dimensions=["eventName", api_name],
                metrics=["eventCount", "activeUsers"],
                order_metric="eventCount",
                limit=300,
            )
        except Exception as exc:
            event_custom_samples[api_name] = [{"error": str(exc)}]

    rows_to_csv(out_dir / "all_events_last_90_days.csv", events)
    rows_to_csv(out_dir / "registered_custom_dimensions.csv", custom_dimensions)
    rows_to_csv(out_dir / "registered_custom_metrics.csv", custom_metrics)
    rows_to_csv(out_dir / "configured_dashboard_definition_validation.csv", validation_rows)
    rows_to_csv(out_dir / "all_metadata_dimensions.csv", all_dimensions)
    rows_to_csv(out_dir / "all_metadata_metrics.csv", all_metrics)
    for api_name, rows in custom_dimension_samples.items():
        safe_name = api_name.replace(":", "_").replace("/", "_")
        rows_to_csv(out_dir / "custom_dimension_value_samples" / f"{safe_name}.csv", rows)
    for api_name, rows in event_custom_samples.items():
        safe_name = api_name.replace(":", "_").replace("/", "_")
        rows_to_csv(out_dir / "event_by_custom_dimension_samples" / f"{safe_name}.csv", rows)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    doc.add_heading("GA4 Events, Custom Parameters & Custom Metrics Audit", 0)
    doc.add_paragraph(f"Property ID: {property_id}")
    doc.add_paragraph("Date range queried: last 90 days ending yesterday.")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(f"Total event names returned by GA4 Data API: {len(events)}")
    doc.add_paragraph(f"Registered custom dimensions exposed by GA4 Data API: {len(custom_dimensions)}")
    doc.add_paragraph(f"Registered custom metrics exposed by GA4 Data API: {len(custom_metrics)}")
    doc.add_paragraph(
        "Important limitation: GA4 Data API exposes event names, standard dimensions/metrics, and registered custom definitions. "
        "It does not list every raw unregistered event parameter. To audit 100% of raw event parameters, enable/use GA4 BigQuery export."
    )

    doc.add_heading("Dashboard Definitions Validation", level=1)
    add_table(doc, validation_rows, ["type", "api_name", "purpose", "status", "sample_rows"], max_rows=20)

    doc.add_heading("Top GA4 Events", level=1)
    add_table(doc, events, ["eventName", "eventCount", "totalUsers", "activeUsers"], max_rows=60)

    doc.add_heading("Registered Custom Dimensions / Parameters", level=1)
    if custom_dimensions:
        add_table(doc, custom_dimensions, ["api_name", "ui_name", "category", "description"], max_rows=80)
    else:
        doc.add_paragraph("No registered custom dimensions were returned by GA4 metadata.")

    doc.add_heading("Registered Custom Metrics", level=1)
    if custom_metrics:
        add_table(doc, custom_metrics, ["api_name", "ui_name", "category", "type", "description"], max_rows=80)
    else:
        doc.add_paragraph("No registered custom metrics were returned by GA4 metadata.")

    doc.add_heading("Sample Values for Custom Dimensions", level=1)
    for api_name, rows in list(custom_dimension_samples.items())[:12]:
        doc.add_heading(api_name, level=2)
        add_table(doc, rows, list(rows[0].keys()) if rows else ["empty"], max_rows=15)

    doc.add_heading("Recommended Dashboard Improvements", level=1)
    recommendations = [
        "Keep Appsmith dashboards reading only from the backend so GA4 credentials stay private.",
        "Use registered custom dimensions for level_number, step_number, and action_type because these power filters and funnels.",
        "For manager-requested exact user journeys such as installed -> never played -> not uninstalled, use BigQuery export because GA4 Data API is aggregate-oriented.",
        "For tutorial funnel, use tutorial step custom dimension plus event counts/users to show where users drop.",
        "For gameplay balancing, keep level performance table with started, completed, drop-off, hints, and average time.",
        "For retention, show both percentage and user count so 0% is understood correctly when cohort age is too recent or cohort users are small.",
    ]
    for rec in recommendations:
        doc.add_paragraph(rec, style="List Bullet")

    docx_path = out_dir / "GA4_Events_Custom_Parameters_Audit.docx"
    doc.save(docx_path)

    md_lines = [
        "# GA4 Events, Custom Parameters & Custom Metrics Audit",
        "",
        f"- Property ID: `{property_id}`",
        "- Date range: last 90 days ending yesterday",
        f"- Event names returned: {len(events)}",
        f"- Registered custom dimensions: {len(custom_dimensions)}",
        f"- Registered custom metrics: {len(custom_metrics)}",
        "",
        "Note: GA4 Data API cannot list every raw unregistered event parameter. BigQuery export is required for a full raw-parameter audit.",
        "",
        "## Files",
        "",
        f"- Word report: `{docx_path.name}`",
        "- CSV: `all_events_last_90_days.csv`",
        "- CSV: `registered_custom_dimensions.csv`",
        "- CSV: `registered_custom_metrics.csv`",
        "- CSV: `configured_dashboard_definition_validation.csv`",
    ]
    (out_dir / "README.md").write_text("\n".join(md_lines) + "\n", encoding="utf-8")

    print(json.dumps({
        "out_dir": str(out_dir),
        "docx": str(docx_path),
        "events": len(events),
        "custom_dimensions": len(custom_dimensions),
        "custom_metrics": len(custom_metrics),
    }, indent=2))


if __name__ == "__main__":
    main()
